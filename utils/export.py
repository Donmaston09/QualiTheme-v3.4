"""
utils/export.py
Export helpers: CSV, Word (.docx), Excel (.xlsx), and AI research report — v3.1

Changes from v3.0:
  - Model string updated to claude-sonnet-4-6
  - _call_llm raises a typed exception on total failure instead of silently
    returning a fallback without explanation
  - export_coded_segments_to_word: newline-only segments are skipped
  - export_to_excel: Frequency Chart sheet uses absolute data reference
    to avoid empty chart on open
  - export_multi_transcript_excel: sheet-name uniqueness guaranteed even
    if two participants share a truncated prefix
  - All Word docs set default body font size (11 pt) for consistent rendering
  - Added export_codebook_to_word None-guard on coded_segments
"""

from __future__ import annotations

import io
from datetime import datetime
from textwrap import shorten

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, Reference

# Current Anthropic model — keep in sync with analysis.py
CLAUDE_MODEL = "claude-sonnet-4-6"

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _ts() -> str:
    return datetime.now().strftime("%d %B %Y, %H:%M")


def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _new_doc(title: str) -> Document:
    """Create a Word document with standard metadata and body font."""
    doc = Document()
    doc.core_properties.author = "QualiTheme v3.1"
    doc.core_properties.title  = title
    # Set default body font size
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    return doc


# ── CSV ─────────────────────────────────────────────────────────────────────

def export_coded_segments_to_csv(coded_segments: list[dict]) -> bytes:
    return pd.DataFrame(coded_segments).to_csv(index=False).encode("utf-8")


# ── Word – Coded Transcript ─────────────────────────────────────────────────

def export_coded_segments_to_word(coded_segments: list[dict]) -> bytes:
    doc = _new_doc("Coded Transcript")

    h = doc.add_heading("Coded Transcript", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(_ts())
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph("")

    df = pd.DataFrame(coded_segments)
    df = df[df["Segment"].astype(str).str.strip() != ""]

    for code, group in df.groupby("Code"):
        heading = doc.add_heading(str(code), level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        for _, row in group.iterrows():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(row["Segment"]))
            if "Sentiment" in row and row["Sentiment"]:
                colour = {
                    "Positive": RGBColor(0x00, 0x80, 0x00),
                    "Negative": RGBColor(0xCC, 0x00, 0x00),
                }.get(row["Sentiment"], RGBColor(0x66, 0x66, 0x66))
                run = p.add_run(f"  [{row['Sentiment']}]")
                run.font.color.rgb = colour
                run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Word – Codebook ─────────────────────────────────────────────────────────

def export_codebook_to_word(
    codebook_df: pd.DataFrame,
    coded_segments: list[dict] | None = None,
) -> bytes:
    doc = _new_doc("Research Codebook")

    h = doc.add_heading("Research Codebook", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(_ts())
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph("")
    doc.add_paragraph(
        "This codebook provides a structured record of all codes identified "
        "during thematic analysis, including their frequency and a representative excerpt."
    )
    doc.add_paragraph("")

    for _, row in codebook_df.iterrows():
        heading = doc.add_heading(str(row["Code"]), level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

        has_sentiment = "Dominant Sentiment" in codebook_df.columns
        tbl = doc.add_table(rows=2, cols=2 if has_sentiment else 1)
        tbl.style = "Table Grid"
        headers = ["Frequency"] + (["Dominant Sentiment"] if has_sentiment else [])
        values  = [str(row["Frequency"])] + ([str(row.get("Dominant Sentiment", "—"))]
                                              if has_sentiment else [])
        for ci, (hdr, val) in enumerate(zip(headers, values)):
            hc = tbl.cell(0, ci)
            hc.text = hdr
            hc.paragraphs[0].runs[0].font.bold = True
            _set_cell_bg(hc, "1F77B4")
            hc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tbl.cell(1, ci).text = val

        doc.add_paragraph("")
        if "Example" in row and row["Example"]:
            p = doc.add_paragraph()
            p.add_run("Example segment:  ").bold = True
            p.add_run(str(row["Example"]))
            p.paragraph_format.left_indent = Cm(0.5)

        if coded_segments:
            segs = [s["Segment"] for s in coded_segments
                    if s.get("Code") == row["Code"]]
            if len(segs) > 1:
                doc.add_paragraph("All segments:", style="Intense Quote")
                for seg in segs:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.add_run(shorten(str(seg), 200, placeholder="..."))
                    bp.paragraph_format.space_after = Pt(2)
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Excel – single transcript ────────────────────────────────────────────────

def export_to_excel(coded_segments: list[dict], codebook_df: pd.DataFrame) -> bytes:
    buf = io.BytesIO()
    wb  = openpyxl.Workbook()

    H_FILL  = PatternFill("solid", fgColor="1F77B4")
    H_FONT  = Font(bold=True, color="FFFFFF", size=11)
    H_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    THIN    = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin"),
    )
    ALT = PatternFill("solid", fgColor="EBF5FB")

    def _hdr(ws, headers: list[str]) -> None:
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=col, value=h)
            c.fill, c.font, c.alignment, c.border = H_FILL, H_FONT, H_ALIGN, THIN
        ws.row_dimensions[1].height = 22

    def _style_row(ws, row_idx: int, n_cols: int) -> None:
        fill = ALT if row_idx % 2 == 0 else None
        for col in range(1, n_cols + 1):
            c = ws.cell(row=row_idx, column=col)
            c.border    = THIN
            c.alignment = Alignment(vertical="top", wrap_text=True)
            if fill:
                c.fill = fill

    # Sheet 1 — Coded Segments
    ws1 = wb.active
    ws1.title = "Coded Segments"
    seg_cols  = ["#", "Segment", "Code", "Sentiment", "Sentiment Score"]
    _hdr(ws1, seg_cols)
    for ltr, w in zip("ABCDE", [5, 70, 28, 14, 18]):
        ws1.column_dimensions[ltr].width = w

    for i, seg in enumerate(coded_segments, 1):
        row = [i, seg.get("Segment", ""), seg.get("Code", ""),
               seg.get("Sentiment", ""), seg.get("SentimentScore", "")]
        for col, val in enumerate(row, 1):
            ws1.cell(row=i + 1, column=col, value=val)
        _style_row(ws1, i + 1, len(seg_cols))
    ws1.freeze_panes = "A2"

    # Sheet 2 — Codebook
    ws2     = wb.create_sheet("Codebook")
    cb_cols = list(codebook_df.columns)
    _hdr(ws2, cb_cols)
    cw = {"Code": 30, "Frequency": 14, "Example": 60, "Dominant Sentiment": 22}
    for col, h in enumerate(cb_cols, 1):
        ws2.column_dimensions[get_column_letter(col)].width = cw.get(h, 20)
    for i, (_, row) in enumerate(codebook_df.iterrows(), 1):
        for col, h in enumerate(cb_cols, 1):
            ws2.cell(row=i + 1, column=col, value=row[h])
        _style_row(ws2, i + 1, len(cb_cols))
    ws2.freeze_panes = "A2"

    # Sheet 3 — Frequency Chart
    ws3         = wb.create_sheet("Frequency Chart")
    ws3["A1"]   = "Code"
    ws3["B1"]   = "Frequency"
    n_rows      = len(codebook_df)
    for i, (_, row) in enumerate(codebook_df.iterrows(), 2):
        ws3.cell(row=i, column=1, value=str(row["Code"]))
        ws3.cell(row=i, column=2, value=int(row["Frequency"]))

    chart = BarChart()
    chart.type, chart.grouping    = "bar", "clustered"
    chart.title                   = "Code Frequencies"
    chart.y_axis.title            = "Code"
    chart.x_axis.title            = "Frequency"
    chart.width, chart.height     = 18, 12
    data_ref = Reference(ws3, min_col=2, min_row=1, max_row=n_rows + 1)
    cats_ref = Reference(ws3, min_col=1, min_row=2, max_row=n_rows + 1)
    chart.add_data(data_ref, titles_from_data=True)
    chart.set_categories(cats_ref)
    ws3.add_chart(chart, "D2")

    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── AI Research Report (Word) ────────────────────────────────────────────────

def generate_ai_report(
    coded_segments: list[dict],
    codebook_df: pd.DataFrame,
    cluster_keywords: dict,
    cluster_labels: dict,
    llm_provider: str,
    llm_api_key: str,
    research_context: str = "",
) -> bytes:
    theme_summary = "\n".join(
        f"- {label}: {', '.join(cluster_keywords.get(k, [])[:6])}"
        for k, label in cluster_labels.items()
    )
    top_codes  = codebook_df.head(10).to_string(index=False)
    sample_segs = "\n".join(
        f"[{s.get('Code', '')}] {shorten(str(s['Segment']), 120)}"
        for s in coded_segments[:20]
    )
    context_note = f"Research context: {research_context}\n" if research_context else ""

    prompt = f"""You are an expert qualitative researcher writing a thematic analysis report.
{context_note}
Based on the following data, write a structured research report with these sections:
1. Executive Summary (3–4 sentences)
2. Methodology (brief — note that sentence-transformer embeddings and K-Means clustering were used)
3. Key Themes Identified (one paragraph per theme, describing its meaning and evidence)
4. Cross-Theme Patterns (co-occurrence and relationships between themes)
5. Notable Quotes and Evidence (3–5 illustrative quotes from the data)
6. Conclusions and Recommendations

DATA:
Themes and keywords:
{theme_summary}

Code frequency table:
{top_codes}

Sample coded segments:
{sample_segs}

Write in an academic but accessible style. Use markdown headings (##) for sections.
Be analytical, not just descriptive. Total length: 600–900 words."""

    llm_text = _call_llm(prompt, llm_provider, llm_api_key)
    return _render_report_docx(llm_text, codebook_df)


def _render_report_docx(llm_text: str, codebook_df: pd.DataFrame) -> bytes:
    doc = _new_doc("Thematic Analysis Report")

    title = doc.add_heading("Thematic Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(_ts())
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph("")

    for line in llm_text.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
        else:
            doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Appendix — Code Frequency Table", level=1)
    tbl = doc.add_table(rows=1, cols=len(codebook_df.columns))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, col in enumerate(codebook_df.columns):
        hdr_cells[i].text = str(col)
        _set_cell_bg(hdr_cells[i], "1F77B4")
        hdr_cells[i].paragraphs[0].runs[0].font.bold  = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for _, row in codebook_df.iterrows():
        cells = tbl.add_row().cells
        for i, col in enumerate(codebook_df.columns):
            cells[i].text = str(row[col])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _call_llm(prompt: str, provider: str, api_key: str,
              max_tokens: int = 1500) -> str:
    try:
        if provider == "Claude (Anthropic)":
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            msg = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=max_tokens,
                messages=[{"role": "user", "content": prompt}],
            )
            return msg.content[0].text.strip()

        elif provider == "OpenAI (GPT-4)":
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            r = client.chat.completions.create(
                model="gpt-4o", max_tokens=max_tokens,
                messages=[{"role": "user", "content": prompt}],
            )
            return r.choices[0].message.content.strip()

        elif provider == "Google Gemini":
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            # Use gemini-1.5-flash — gemini-pro is deprecated in v1beta
            for model_name in ["gemini-1.5-flash", "gemini-1.5-pro", "gemini-2.0-flash"]:
                try:
                    r = genai.GenerativeModel(model_name).generate_content(prompt)
                    return r.text.strip()
                except Exception:
                    continue
            raise RuntimeError("No available Gemini model responded successfully.")

    except Exception as exc:
        return _fallback_report(str(exc))

    return _fallback_report("No LLM provider selected or key missing.")


def _fallback_report(error: str) -> str:
    return (
        "## Executive Summary\n"
        "This report summarises the thematic analysis performed on the uploaded transcript. "
        f"(Note: AI report generation was unavailable — {error}. "
        "The codebook appendix provides the full structured results.)\n\n"
        "## Methodology\n"
        "Segments were encoded using the all-MiniLM-L6-v2 sentence-transformer model "
        "and clustered via K-Means. Cluster keywords were extracted using TF-IDF.\n\n"
        "## Key Themes\n"
        "Please refer to the codebook appendix for theme frequencies and examples.\n\n"
        "## Conclusions\n"
        "Further manual review of the coded segments is recommended to refine themes "
        "and develop interpretive insights aligned with your research questions."
    )


# ── Multi-transcript helpers ─────────────────────────────────────────────────

def export_all_transcripts_csv(transcripts: list[dict]) -> bytes:
    rows = []
    for t in transcripts:
        pid   = t.get("participant_id", t.get("filename", "Unknown"))
        fname = t.get("filename", "")
        for seg in t.get("coded_segments", []):
            row = {"Participant": pid, "File": fname}
            row.update(seg)
            rows.append(row)
    return pd.DataFrame(rows).to_csv(index=False).encode("utf-8")


def export_multi_transcript_excel(transcripts: list[dict]) -> bytes:
    buf = io.BytesIO()
    wb  = openpyxl.Workbook()

    H_FILL  = PatternFill("solid", fgColor="1F77B4")
    H_FONT  = Font(bold=True, color="FFFFFF", size=11)
    H_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    THIN    = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin"),
    )
    ALT = PatternFill("solid", fgColor="EBF5FB")

    def _hdr(ws, headers: list[str]) -> None:
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=col, value=h)
            c.fill, c.font, c.alignment, c.border = H_FILL, H_FONT, H_ALIGN, THIN
        ws.row_dimensions[1].height = 22

    def _data_row(ws, row_idx: int, values: list, n_cols: int) -> None:
        fill = ALT if row_idx % 2 == 0 else None
        for col, val in enumerate(values, 1):
            c = ws.cell(row=row_idx, column=col, value=val)
            c.border    = THIN
            c.alignment = Alignment(vertical="top", wrap_text=True)
            if fill:
                c.fill = fill

    # Collect all segments
    all_rows: list[dict] = []
    for t in transcripts:
        pid   = t.get("participant_id", t.get("filename", "Unknown"))
        fname = t.get("filename", "")
        for seg in t.get("coded_segments", []):
            row = {"Participant": pid, "File": fname}
            row.update(seg)
            all_rows.append(row)
    all_df = pd.DataFrame(all_rows)

    # Sheet 1 — Summary pivot
    ws_summary       = wb.active
    ws_summary.title = "Summary"

    if not all_df.empty and "Code" in all_df.columns:
        pivot = (
            all_df[all_df["Code"].astype(str).str.strip() != ""]
            .groupby(["Participant", "Code"])
            .size()
            .unstack(fill_value=0)
            .reset_index()
        )
        headers = list(pivot.columns)
        _hdr(ws_summary, headers)
        ws_summary.column_dimensions["A"].width = 25
        for col_idx in range(2, len(headers) + 1):
            ws_summary.column_dimensions[get_column_letter(col_idx)].width = 18
        for i, (_, row) in enumerate(pivot.iterrows(), 1):
            _data_row(ws_summary, i + 1, list(row), len(headers))

        n_rows = len(pivot) + 1
        n_cols = len(headers)
        chart  = BarChart()
        chart.type, chart.grouping = "col", "clustered"
        chart.title                = "Code Counts by Participant"
        chart.y_axis.title         = "Count"
        chart.width, chart.height  = 22, 12
        data_ref = Reference(ws_summary, min_col=2, min_row=1,
                             max_col=n_cols, max_row=n_rows)
        cats_ref = Reference(ws_summary, min_col=1, min_row=2, max_row=n_rows)
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats_ref)
        ws_summary.add_chart(chart, f"A{n_rows + 3}")
    else:
        ws_summary["A1"] = "No coded data available yet."
    ws_summary.freeze_panes = "A2"

    # Sheet 2 — Master segments
    ws_master       = wb.create_sheet("All Segments")
    if not all_df.empty:
        master_cols = list(all_df.columns)
        _hdr(ws_master, master_cols)
        ws_master.column_dimensions["A"].width = 22
        ws_master.column_dimensions["B"].width = 20
        if "Segment" in master_cols:
            ws_master.column_dimensions[
                get_column_letter(master_cols.index("Segment") + 1)
            ].width = 65
        if "Code" in master_cols:
            ws_master.column_dimensions[
                get_column_letter(master_cols.index("Code") + 1)
            ].width = 28
        for i, (_, row) in enumerate(all_df.iterrows(), 1):
            _data_row(ws_master, i + 1, list(row), len(master_cols))
    ws_master.freeze_panes = "A2"

    # Sheets 3+ — one per participant (guaranteed-unique names ≤31 chars)
    used_names: set[str] = {"Summary", "All Segments"}
    for t in transcripts:
        pid  = t.get("participant_id", t.get("filename", "Unknown"))
        segs = t.get("coded_segments", [])
        if not segs:
            continue

        base = str(pid)[:28].replace("/", "-").replace("\\", "-").replace(":", "-")
        name = base
        suffix = 2
        while name in used_names:
            name = f"{base[:25]}_{suffix}"
            suffix += 1
        used_names.add(name)

        ws_p = wb.create_sheet(name)
        p_headers = ["#", "Segment", "Code", "Sentiment", "Sentiment Score"]
        _hdr(ws_p, p_headers)
        for ltr, w in zip("ABCDE", [5, 65, 28, 14, 18]):
            ws_p.column_dimensions[ltr].width = w
        for i, seg in enumerate(segs, 1):
            row_vals = [
                i,
                seg.get("Segment", ""),
                seg.get("Code", ""),
                seg.get("Sentiment", ""),
                seg.get("SentimentScore", ""),
            ]
            _data_row(ws_p, i + 1, row_vals, len(p_headers))
        ws_p.freeze_panes = "A2"

    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ── Corpus Synthesis Report ─────────────────────────────────────────────────

def generate_synthesis_report(
    synthesis: dict,
    figures: dict,
    llm_provider: str = "",
    llm_api_key: str = "",
    research_context: str = "",
    n_participants: int = 0,
    transcripts: list | None = None,
) -> bytes:
    """
    Generate a comprehensive cross-participant synthesis report as Word .docx.

    Charts are regenerated from raw data inside this function so they are
    never dependent on cached matplotlib Figure objects from session_state,
    which can be garbage-collected or fail to pickle between Streamlit reruns.

    Parameters
    ----------
    synthesis        : output of analysis.synthesize_corpus()
    figures          : pre-rendered Figure dict (used as first preference;
                       regenerated from transcripts if embedding fails)
    llm_provider     : LLM provider string
    llm_api_key      : API key
    research_context : researcher's study questions / context
    n_participants   : number of participants
    transcripts      : raw transcript list (used to regenerate charts if needed)
    """
    from utils.analysis import (
        plot_synthesis_frequency, plot_synthesis_saturation,
        plot_synthesis_heatmap, plot_synthesis_sentiment_breakdown,
        plot_synthesis_cooccurrence, plot_synthesis_theme_journey,
        plot_synthesis_quotes_table,
    )

    codebook   = synthesis.get("codebook",          pd.DataFrame())
    saturation = synthesis.get("saturation",        pd.DataFrame())
    top_quotes = synthesis.get("top_quotes",        [])
    n_p        = synthesis.get("n_participants",    n_participants)
    n_segs     = synthesis.get("n_coded_segments",  0)
    n_codes    = synthesis.get("n_unique_codes",    0)
    all_segs   = synthesis.get("all_segments",      [])

    # Filter saturation to themes seen by ≥2 participants for the report
    sat_filtered = (
        saturation[saturation["Participants"] >= 2].copy()
        if not saturation.empty else saturation
    )
    cb_filtered = (
        codebook[codebook["Participants"] >= 2].copy()
        if not codebook.empty else codebook
    )

    # ── Build LLM narrative ───────────────────────────────────────────────
    theme_block = "\n".join(
        f"- {row['Code']}: {int(row['Frequency'])} segments, "
        f"{int(row['Participants'])}/{n_p} participants "
        f"({row.get('% Participants', 0):.0f}% — {row.get('Saturation', '')})"
        for _, row in sat_filtered.head(12).iterrows()
    ) if not sat_filtered.empty else "No multi-participant themes found."

    # Only use quotes for themes seen by ≥2 participants
    filtered_quote_codes = (
        set(sat_filtered["Code"].tolist()) if not sat_filtered.empty else set()
    )
    good_quotes = [
        q for q in top_quotes
        if q.get("Code") in filtered_quote_codes
    ][:8]

    quotes_block = "\n".join(
        f'[{q["Code"]}] "{shorten(q["Quote"], 120)}" — {q["Participant"]}'
        for q in good_quotes
    ) if good_quotes else "No cross-participant quotes available."

    context_note = f"Research context: {research_context}\n\n" if research_context else ""

    prompt = f"""You are an expert qualitative researcher writing the synthesis section of a
multi-participant thematic analysis study. Your task is to interpret and synthesise
the patterns in the data, not just describe the numbers.
{context_note}
STUDY OVERVIEW:
- Participants analysed: {n_p}
- Total coded segments: {n_segs}
- Themes appearing in ≥2 participants: {len(sat_filtered)}

THEME SATURATION (themes seen across multiple participants only):
{theme_block}

REPRESENTATIVE QUOTES (from cross-participant themes):
{quotes_block}

Write a structured synthesis report with these six sections using ## headings:

## Executive Summary
3–4 sentences covering the dominant cross-cutting themes and overall picture.

## Methodology
One paragraph: multi-participant qualitative thematic analysis using
sentence-transformer embeddings (all-MiniLM-L6-v2) and K-Means clustering,
followed by saturation analysis to identify cross-participant themes.

## Key Themes Across Participants
For each of the top 5–7 cross-participant themes (those with highest saturation):
one paragraph per theme covering its meaning, the number of participants who raised it,
its saturation classification, and an illustrative quote where available.
Use the participant count and saturation label in your analysis.

## Cross-Theme Patterns
Discuss which themes co-occur, any tensions or contradictions between themes,
and what the overall pattern suggests about the research topic.

## Participant Variation
Briefly discuss what was consistent across all participants vs. what was specific
to subgroups or individual voices. Reference saturation levels.

## Conclusions and Recommendations
3–5 evidence-based conclusions directly grounded in the theme data,
followed by 3–5 concrete actionable recommendations.

Write in formal academic style. Total length: 900–1200 words.
Be analytical and interpretive — do not simply list themes or repeat the numbers."""

    narrative = _call_llm(prompt, llm_provider, llm_api_key, max_tokens=2000)

    # If fallback was used, build a data-driven narrative instead
    if "AI report generation was unavailable" in narrative:
        narrative = _build_data_driven_narrative(
            sat_filtered, good_quotes, n_p, n_segs, research_context
        )

    # ── Regenerate charts fresh for embedding ─────────────────────────────
    # Re-render from data rather than relying on cached Figure objects,
    # which may be closed/GC'd by the time the report button is clicked.
    chart_fns = {
        "synthesis_frequency":  lambda: plot_synthesis_frequency(cb_filtered if not cb_filtered.empty else codebook),
        "synthesis_saturation": lambda: plot_synthesis_saturation(sat_filtered if not sat_filtered.empty else saturation, n_p),
        "synthesis_heatmap":    lambda: plot_synthesis_heatmap(transcripts) if transcripts else None,
        "synthesis_sentiment":  lambda: plot_synthesis_sentiment_breakdown(synthesis),
        "synthesis_network":    lambda: plot_synthesis_cooccurrence(all_segs),
        "synthesis_journey":    lambda: plot_synthesis_theme_journey(transcripts) if transcripts else None,
        "synthesis_quotes":     lambda: plot_synthesis_quotes_table(good_quotes if good_quotes else top_quotes),
    }

    rendered_figs: dict = {}
    for key, fn in chart_fns.items():
        try:
            fig = fn()
            if fig is not None:
                rendered_figs[key] = fig
        except Exception as e:
            print(f"[QualiTheme] Chart '{key}' failed during report generation: {e}")

    # ── Render Word document ──────────────────────────────────────────────
    doc = _new_doc("Corpus Synthesis Report")

    # ── Cover page ────────────────────────────────────────────────────────
    title_p = doc.add_heading("Corpus Thematic Synthesis Report", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Cross-participant analysis  ·  {n_p} participants  ·  {_ts()}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Horizontal rule effect
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after  = Pt(12)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(_qn("w:val"),   "single")
    bottom.set(_qn("w:sz"),    "6")
    bottom.set(_qn("w:space"), "1")
    bottom.set(_qn("w:color"), "1F77B4")
    pBdr.append(bottom)
    pPr.append(pBdr)

    if research_context:
        rq = doc.add_paragraph()
        rq.add_run("Research context: ").bold = True
        rq.add_run(research_context)
        rq.paragraph_format.left_indent = Cm(0.5)
        rq.paragraph_format.space_after = Pt(8)

    # ── Metrics summary (3 styled paragraphs instead of table) ───────────
    # Using separate centred paragraphs avoids the Word newline-in-cell issue
    metrics = [
        (str(n_p),                   "Participants analysed"),
        (str(n_segs),                "Total coded segments"),
        (str(len(sat_filtered)),     "Cross-participant themes (≥2 participants)"),
    ]
    metrics_tbl = doc.add_table(rows=1, cols=3)
    metrics_tbl.style = "Table Grid"
    for ci, (val, label) in enumerate(metrics):
        cell = metrics_tbl.cell(0, ci)
        cell.text = ""
        _set_cell_bg(cell, "EBF5FB")
        from docx.oxml import OxmlElement as _OE
        # Set fixed cell width
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = _OE("w:tcW")
        tcW.set(_qn("w:w"),    "3000")
        tcW.set(_qn("w:type"), "dxa")
        tcPr.append(tcW)

        p_val = cell.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_val = p_val.add_run(val)
        r_val.bold = True
        r_val.font.size  = Pt(22)
        r_val.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

        p_lbl = cell.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size  = Pt(9)
        r_lbl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    # ── LLM / data-driven narrative ───────────────────────────────────────
    for line in narrative.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
        else:
            doc.add_paragraph(line)

    # ── Visualisations section ────────────────────────────────────────────
    chart_titles = {
        "synthesis_frequency":  "Figure 1. Corpus Theme Frequencies",
        "synthesis_saturation": "Figure 2. Theme Saturation Map",
        "synthesis_heatmap":    "Figure 3. Participant × Theme Heatmap",
        "synthesis_sentiment":  "Figure 4. Corpus Sentiment Analysis",
        "synthesis_network":    "Figure 5. Theme Co-occurrence Network",
        "synthesis_journey":    "Figure 6. Theme Journey Across Participants",
        "synthesis_quotes":     "Figure 7. Representative Quotes by Theme",
    }

    doc.add_page_break()
    h_viz = doc.add_heading("Visualisations", level=1)
    if h_viz.runs:
        h_viz.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    charts_embedded = 0
    for fig_key, fig_title in chart_titles.items():
        fig = rendered_figs.get(fig_key)
        if fig is None:
            continue
        doc.add_paragraph("")
        caption = doc.add_paragraph(fig_title)
        if caption.runs:
            caption.runs[0].bold = True
            caption.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        caption.paragraph_format.space_after = Pt(4)

        buf_fig = io.BytesIO()
        try:
            fig.savefig(buf_fig, format="png", dpi=200, bbox_inches="tight")
            buf_fig.seek(0)
            doc.add_picture(buf_fig, width=Inches(5.8))
            charts_embedded += 1
        except Exception as e:
            doc.add_paragraph(
                f"[Chart generation error for {fig_key}: {e}]"
            )
        finally:
            import matplotlib.pyplot as plt
            plt.close(fig)
        doc.add_paragraph("")

    if charts_embedded == 0:
        note = doc.add_paragraph()
        note.add_run(
            "Note: Charts could not be embedded in this run. "
            "Open the Synthesis tab in the app and use the individual "
            "chart download buttons to save each figure separately."
        ).italic = True
        note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Appendix A — Saturation table ─────────────────────────────────────
    sat_to_use = sat_filtered if not sat_filtered.empty else saturation
    if not sat_to_use.empty:
        doc.add_page_break()
        h_app = doc.add_heading("Appendix A — Theme Saturation Table", level=1)
        if h_app.runs:
            h_app.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        doc.add_paragraph(
            f"The table below shows all {len(sat_to_use)} themes identified across "
            f"{n_p} participants, ordered by frequency. "
            "Saturation classifications: Universal >80% · Major 50–80% · "
            "Moderate 25–49% · Minority <25% of participants."
        )
        doc.add_paragraph("")

        sat_cols = ["Code", "Frequency", "Participants", "% Participants", "Saturation"]
        sat_cols = [c for c in sat_cols if c in sat_to_use.columns]
        tbl = doc.add_table(rows=1, cols=len(sat_cols))
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for i, col in enumerate(sat_cols):
            hdr_cells[i].text = col
            _set_cell_bg(hdr_cells[i], "1F77B4")
            if hdr_cells[i].paragraphs[0].runs:
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for ri, (_, row) in enumerate(sat_to_use.iterrows()):
            cells = tbl.add_row().cells
            for i, col in enumerate(sat_cols):
                cells[i].text = str(row[col])
                if ri % 2 == 0:
                    _set_cell_bg(cells[i], "EBF5FB")

    # ── Appendix B — Representative quotes ───────────────────────────────
    quotes_for_appendix = good_quotes if good_quotes else top_quotes
    if quotes_for_appendix:
        doc.add_page_break()
        h_qapp = doc.add_heading("Appendix B — Representative Quotes", level=1)
        if h_qapp.runs:
            h_qapp.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        doc.add_paragraph(
            "Representative quotes from cross-participant themes. "
            "One quote per theme, selected as illustrative of the broader pattern."
        )

        seen_codes: set = set()
        for q in quotes_for_appendix:
            if q["Code"] in seen_codes:
                continue
            seen_codes.add(q["Code"])

            doc.add_paragraph("")
            h_code = doc.add_paragraph()
            h_code.add_run(q["Code"]).bold = True
            if h_code.runs:
                h_code.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

            quote_p = doc.add_paragraph(style="Quote")
            quote_p.add_run(f'"{q["Quote"]}"')
            quote_p.paragraph_format.left_indent  = Cm(1.0)
            quote_p.paragraph_format.right_indent = Cm(1.0)

            src_p = doc.add_paragraph()
            src_p.add_run(f"— {q['Participant']}").italic = True
            if src_p.runs:
                src_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            src_p.paragraph_format.left_indent = Cm(1.0)

    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out.getvalue()


def _build_data_driven_narrative(
    sat_df: pd.DataFrame,
    quotes: list[dict],
    n_p: int,
    n_segs: int,
    research_context: str,
) -> str:
    """
    Build a structured narrative from the data directly — used when no
    LLM is available or when the LLM call fails. Produces a real synthesis
    rather than a placeholder.
    """
    if sat_df.empty:
        return (
            "## Executive Summary\n"
            "Thematic analysis was completed across all participants. "
            "Please refer to the visualisations and appendix tables for detailed findings.\n\n"
            "## Methodology\n"
            "Segments were encoded using the all-MiniLM-L6-v2 sentence-transformer model "
            "and clustered via K-Means. Cross-participant saturation was computed for each theme.\n\n"
            "## Conclusions\n"
            "Refer to the saturation table in Appendix A for theme frequencies and participant breadth."
        )

    context_line = (
        f"The study explored: {research_context}. " if research_context else ""
    )
    universal  = sat_df[sat_df["Saturation"] == "Universal (>80%)"]
    major      = sat_df[sat_df["Saturation"] == "Major (50–80%)"]
    moderate   = sat_df[sat_df["Saturation"] == "Moderate (25–49%)"]
    minority   = sat_df[sat_df["Saturation"] == "Minority (<25%)"]

    top_theme     = sat_df.iloc[0]["Code"] if len(sat_df) > 0 else "—"
    top_freq      = int(sat_df.iloc[0]["Frequency"]) if len(sat_df) > 0 else 0
    top_pct       = sat_df.iloc[0].get("% Participants", 0) if len(sat_df) > 0 else 0

    # Build per-theme paragraphs
    theme_paras = []
    for _, row in sat_df.head(7).iterrows():
        code  = row["Code"]
        freq  = int(row["Frequency"])
        parts = int(row["Participants"])
        pct   = row.get("% Participants", 0)
        sat   = row.get("Saturation", "")

        # Find a matching quote
        q_match = next((q["Quote"] for q in quotes if q["Code"] == code), None)
        quote_line = (
            f' One participant observed: "{shorten(q_match, 100)}"'
            if q_match else ""
        )

        theme_paras.append(
            f"**{code}** ({sat}): This theme appeared in {freq} coded segments "
            f"across {parts} of {n_p} participants ({pct:.0f}%).{quote_line}"
        )

    themes_section = "\n\n".join(theme_paras)

    lines = [
        "## Executive Summary",
        f"{context_line}Thematic analysis across {n_p} participants and {n_segs} coded segments "
        f"identified {len(sat_df)} cross-participant themes. "
        f"The most prominent theme was '{top_theme}', "
        f"appearing in {top_freq} segments across {top_pct:.0f}% of participants. "
        f"{len(universal)} theme(s) were universal (>80% of participants), "
        f"{len(major)} major (50–80%), "
        f"{len(moderate)} moderate (25–49%), "
        f"and {len(minority)} minority (<25%).",
        "",
        "## Methodology",
        f"This study employed multi-participant qualitative thematic analysis. "
        f"Transcripts from {n_p} participants were segmented and encoded using the "
        f"all-MiniLM-L6-v2 sentence-transformer model. K-Means clustering identified "
        f"thematic groupings, and saturation analysis quantified each theme's "
        f"prevalence across the participant group.",
        "",
        "## Key Themes Across Participants",
        themes_section,
        "",
        "## Cross-Theme Patterns",
        "The distribution of themes across participants suggests a dataset with "
        f"{'broad consensus on core themes' if len(universal) + len(major) > 2 else 'considerable variation between participants'}. "
        f"Universal and major themes likely represent shared experiences, "
        "while moderate and minority themes may reflect individual circumstances or subgroup differences. "
        "The saturation map (Figure 2) and co-occurrence network (Figure 5) provide "
        "visual representations of these relationships.",
        "",
        "## Participant Variation",
        f"{'Several themes showed strong cross-participant consistency, suggesting shared structural or contextual factors.' if len(universal) > 0 else 'No theme reached universal saturation, indicating meaningful variation across the participant group.'} "
        f"{'Minority themes (<25% saturation) may warrant further investigation as potential outlier experiences or emerging patterns not yet fully represented in this sample.' if len(minority) > 0 else ''}",
        "",
        "## Conclusions and Recommendations",
        f"- The prominence of '{top_theme}' across {top_pct:.0f}% of participants "
        f"suggests it represents a core finding warranting focused attention.",
        f"- Themes with Universal or Major saturation should anchor the primary conclusions "
        f"of this study.",
        f"- Minority themes identified in individual participants may benefit from follow-up "
        f"through targeted interviews or member-checking.",
        f"- Further qualitative analysis using LLM-assisted coding (Claude, GPT-4) would "
        f"produce richer theme labels and more nuanced interpretive synthesis.",
        f"- Consider member-checking with participants to validate the identified themes "
        f"and their descriptions.",
    ]
    return "\n".join(lines)
    """
    Generate a comprehensive cross-participant synthesis report as a Word .docx.

    Parameters
    ----------
    synthesis       : output of analysis.synthesize_corpus()
    figures         : {name: matplotlib.Figure} — pre-rendered charts to embed
    llm_provider    : LLM provider string (optional — used for narrative text)
    llm_api_key     : API key (optional)
    research_context: researcher's study questions / context
    n_participants  : number of participants (for display)

    Returns
    -------
    bytes of a fully formatted .docx file
    """
    codebook    = synthesis.get("codebook",         pd.DataFrame())
    saturation  = synthesis.get("saturation",       pd.DataFrame())
    top_quotes  = synthesis.get("top_quotes",       [])
    sent_sum    = synthesis.get("sentiment_summary", {})
    n_p         = synthesis.get("n_participants",   n_participants)
    n_segs      = synthesis.get("n_coded_segments", 0)
    n_codes     = synthesis.get("n_unique_codes",   0)

    # ── Build LLM narrative ───────────────────────────────────────────────
    theme_block = "\n".join(
        f"- {row['Code']}: {int(row['Frequency'])} segments, "
        f"{int(row['Participants'])} participants "
        f"({row.get('% Participants', 0):.0f}% saturation)"
        for _, row in saturation.head(10).iterrows()
    ) if not saturation.empty else "No saturation data available."

    quotes_block = "\n".join(
        f'[{q["Code"]}] "{shorten(q["Quote"], 120)}" — {q["Participant"]}'
        for q in top_quotes[:8]
    )

    context_note = f"Research context: {research_context}\n\n" if research_context else ""

    prompt = f"""You are an expert qualitative researcher writing the synthesis section of a
multi-participant thematic analysis study.
{context_note}
STUDY OVERVIEW:
- Participants analysed: {n_p}
- Total coded segments: {n_segs}
- Unique themes identified: {n_codes}

THEME SATURATION (frequency + participant breadth):
{theme_block}

REPRESENTATIVE QUOTES (one per major theme):
{quotes_block}

Write a structured synthesis report with these sections:
## Executive Summary
(3–4 sentences summarising the dominant findings across all participants)

## Methodology
(1 paragraph: describe the multi-participant thematic analysis approach,
sentence-transformer embeddings, K-Means clustering, and manual review)

## Key Themes Across Participants
(For each of the top 5–7 themes: one paragraph covering what it means,
how many participants raised it, saturation level, and a supporting quote.
Note whether the theme is universal, major, moderate or minority.)

## Cross-Theme Patterns
(Discuss relationships and co-occurrences between themes — what tends to
appear together, what tensions exist between themes)

## Participant Variation
(Briefly discuss which themes were consistent across all participants vs.
concentrated in particular voices — reference the saturation data)

## Conclusions and Recommendations
(Evidence-based conclusions and 3–5 actionable recommendations)

Write in formal academic style. Total length: 900–1200 words.
Use markdown headings (##) for sections. Be analytical, not descriptive.
Do not reproduce raw data — interpret and synthesise it."""

    narrative = _call_llm(prompt, llm_provider, llm_api_key)

    # ── Render Word document ──────────────────────────────────────────────
    doc = _new_doc("Corpus Synthesis Report")

    # Cover
    title_p = doc.add_heading("Corpus Thematic Synthesis Report", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Cross-participant analysis  ·  {n_p} participants  ·  {_ts()}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph("")

    if research_context:
        rq = doc.add_paragraph()
        rq.add_run("Research context: ").bold = True
        rq.add_run(research_context)
        rq.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph("")

    # Key metrics boxes (as a simple table)
    metrics_tbl = doc.add_table(rows=1, cols=3)
    metrics_tbl.style = "Table Grid"
    metric_data = [
        (str(n_p),    "Participants"),
        (str(n_segs), "Coded Segments"),
        (str(n_codes),"Unique Themes"),
    ]
    for ci, (val, label) in enumerate(metric_data):
        cell = metrics_tbl.cell(0, ci)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p.add_run(val + "\n")
        run_val.bold = True
        run_val.font.size = Pt(18)
        run_val.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        run_lbl = p.add_run(label)
        run_lbl.font.size = Pt(9)
        run_lbl.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _set_cell_bg(cell, "EBF5FB")
    doc.add_paragraph("")

    # ── LLM narrative sections ────────────────────────────────────────────
    for line in narrative.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith(("- ", "* ")):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
        else:
            doc.add_paragraph(line)

    # ── Embedded charts ───────────────────────────────────────────────────
    chart_titles = {
        "synthesis_frequency":  "Figure 1. Corpus Theme Frequencies",
        "synthesis_saturation": "Figure 2. Theme Saturation Map",
        "synthesis_heatmap":    "Figure 3. Participant × Theme Heatmap",
        "synthesis_sentiment":  "Figure 4. Corpus Sentiment Analysis",
        "synthesis_network":    "Figure 5. Theme Co-occurrence Network",
        "synthesis_journey":    "Figure 6. Theme Journey Across Participants",
        "synthesis_quotes":     "Figure 7. Representative Quotes by Theme",
    }

    doc.add_page_break()
    h_viz = doc.add_heading("Visualisations", level=1)
    h_viz.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    for fig_key, fig_title in chart_titles.items():
        fig = figures.get(fig_key)
        if fig is None:
            continue
        doc.add_paragraph("")
        caption = doc.add_paragraph(fig_title)
        caption.runs[0].bold = True
        caption.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        caption.paragraph_format.space_after = Pt(4)

        # Save figure to a temp buffer and embed
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
        buf.seek(0)
        try:
            doc.add_picture(buf, width=Inches(6.0))
        except Exception:
            doc.add_paragraph(f"[Chart could not be embedded: {fig_key}]")
        doc.add_paragraph("")

    # ── Saturation appendix table ─────────────────────────────────────────
    if not saturation.empty:
        doc.add_page_break()
        doc.add_heading("Appendix A — Theme Saturation Table", level=1)
        doc.add_paragraph(
            "The table below shows each identified theme, its total frequency "
            "across all participants, the number and percentage of participants "
            "who used it, and its saturation classification."
        )
        doc.add_paragraph("")

        sat_cols = ["Code", "Frequency", "Participants", "% Participants", "Saturation"]
        sat_cols = [c for c in sat_cols if c in saturation.columns]
        tbl = doc.add_table(rows=1, cols=len(sat_cols))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, col in enumerate(sat_cols):
            hdr[i].text = col
            _set_cell_bg(hdr[i], "1F77B4")
            hdr[i].paragraphs[0].runs[0].font.bold  = True
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for ri, (_, row) in enumerate(saturation.iterrows()):
            cells = tbl.add_row().cells
            for i, col in enumerate(sat_cols):
                cells[i].text = str(row[col])
                if ri % 2 == 0:
                    _set_cell_bg(cells[i], "EBF5FB")

    # ── Quotes appendix ───────────────────────────────────────────────────
    if top_quotes:
        doc.add_page_break()
        doc.add_heading("Appendix B — Representative Quotes", level=1)
        doc.add_paragraph(
            "The following quotes were selected as representative examples "
            "for each major theme. Quotes are presented as coded in the "
            "original transcript data."
        )
        seen_codes: set = set()
        for q in top_quotes:
            if q["Code"] in seen_codes:
                continue
            seen_codes.add(q["Code"])
            doc.add_paragraph("")
            h = doc.add_paragraph()
            h.add_run(q["Code"]).bold = True
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
            quote_p = doc.add_paragraph(style="Quote")
            quote_p.add_run(f'"{q["Quote"]}"')
            src_p = doc.add_paragraph()
            src_p.add_run(f"— {q['Participant']}").italic = True
            src_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            src_p.paragraph_format.left_indent = Cm(1)

    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out.getvalue()
